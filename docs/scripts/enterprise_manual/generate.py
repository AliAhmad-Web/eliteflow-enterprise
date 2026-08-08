#!/usr/bin/env python3
"""Generate EliteFlow Product Documentation (PDF + DOCX) — internship product manual."""

from __future__ import annotations

import re
import sys
from pathlib import Path

_PKG = Path(__file__).resolve().parent
_SCRIPTS = _PKG.parent
if str(_SCRIPTS) not in sys.path:
    sys.path.insert(0, str(_SCRIPTS))

from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from playwright.sync_api import sync_playwright

from enterprise_manual import brand, html_content, paths
from enterprise_manual.styles import css

# Equal page margins — header/footer live inside these bands (no content overlap).
MARGIN_MM = "16mm"
HEADER_MM = "14mm"
FOOTER_MM = "14mm"


def build_html() -> str:
    body = html_content.build_body_html()
    return f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="utf-8"/>
  <title>{brand.DOC_TITLE} — Product Documentation</title>
  <style>{css()}</style>
</head>
<body>
{body}
</body>
</html>
"""


def extract_heading_pages(pdf_path: Path) -> dict[str, int]:
    import fitz

    doc = fitz.open(str(pdf_path))
    mapping: dict[str, int] = {}
    for i, page in enumerate(doc):
        blocks = page.get_text("dict")["blocks"]
        for block in blocks:
            for line in block.get("lines", []):
                spans = line.get("spans", [])
                if not spans:
                    continue
                size = max(s.get("size", 0) for s in spans)
                text = "".join(s.get("text", "") for s in spans).strip()
                if size >= 13 and text:
                    mapping.setdefault(text, i + 1)
    return mapping


def apply_toc_pages(html_doc: str, heading_pages: dict[str, int]) -> str:
    soup = BeautifulSoup(html_doc, "lxml")
    for li in soup.select(".toc-list li, .lof-list li, .lot-list li"):
        a = li.find("a")
        ref = li.find(class_="page-ref")
        if ref is None:
            continue
        title = a.get_text(strip=True) if a else li.get_text(strip=True)
        clean = re.sub(r"^(Figure|Table)\s+\d+\s*[—\-–]\s*", "", title).strip()
        page_no = heading_pages.get(title) or heading_pages.get(clean)
        if page_no:
            ref.string = str(page_no)
    return str(soup)


def _header_template() -> str:
    return """
    <div style="width:100%;box-sizing:border-box;padding:0 2mm;font-family:'Segoe UI',Arial,sans-serif;font-size:8px;color:#4B5563;display:flex;justify-content:space-between;align-items:flex-end;border-bottom:1px solid #D1D5DB;padding-bottom:3px;margin:0 8mm;">
      <div>
        <span style="color:#0B3A6E;font-weight:700;">EliteFlow</span>
        <span style="color:#7C3AED;"> · Product Documentation</span>
      </div>
      <div style="color:#6B7280;">Internship Project</div>
    </div>
    """


def _footer_template() -> str:
    return f"""
    <div style="width:100%;box-sizing:border-box;padding:0 2mm;font-family:'Segoe UI',Arial,sans-serif;font-size:8px;color:#4B5563;display:flex;justify-content:space-between;align-items:flex-start;border-top:1px solid #D1D5DB;padding-top:3px;margin:0 8mm;">
      <div>Confidential · EliteFlow</div>
      <div>Page <span class="pageNumber"></span> of <span class="totalPages"></span></div>
      <div>v{brand.DOC_VERSION}</div>
    </div>
    """


def render_pdf_raw(html_doc: str, dest: Path, *, with_chrome: bool = True) -> None:
    paths.ASSETS.mkdir(parents=True, exist_ok=True)
    paths.HTML_PATH.write_text(html_doc, encoding="utf-8")
    with sync_playwright() as p:
        browser = p.chromium.launch()
        page = browser.new_page()
        page.goto(paths.HTML_PATH.as_uri(), wait_until="networkidle")
        page.evaluate(
            """async () => {
              const imgs = Array.from(document.images);
              await Promise.all(imgs.map(img => {
                if (img.complete && img.naturalWidth > 0) return true;
                return new Promise(resolve => {
                  img.onload = () => resolve(true);
                  img.onerror = () => resolve(false);
                });
              }));
            }"""
        )
        page.pdf(
            path=str(dest),
            format="A4",
            print_background=True,
            display_header_footer=with_chrome,
            header_template=_header_template() if with_chrome else "<div></div>",
            footer_template=_footer_template() if with_chrome else "<div></div>",
            margin={
                "top": HEADER_MM if with_chrome else MARGIN_MM,
                "bottom": FOOTER_MM if with_chrome else MARGIN_MM,
                "left": MARGIN_MM,
                "right": MARGIN_MM,
            },
            prefer_css_page_size=True,
        )
        browser.close()


def write_pdf(src: Path, dest: Path) -> Path:
    """Copy/replace PDF; fall back to alternate name if locked."""
    import shutil

    try:
        if dest.exists():
            dest.unlink()
        shutil.copyfile(src, dest)
        print(f"PDF written: {dest}")
        return dest
    except PermissionError:
        alt = dest.with_name(dest.stem + "_PRODUCT.pdf")
        try:
            if alt.exists():
                alt.unlink()
        except OSError:
            pass
        shutil.copyfile(src, alt)
        print(f"WARNING: {dest.name} locked — wrote {alt.name}")
        print(f"PDF written: {alt}")
        return alt


def generate_pdf(html_doc: str) -> Path:
    raw_pdf = paths.ASSETS / "manual-raw.pdf"
    print("Rendering PDF (pass 1)…")
    render_pdf_raw(html_doc, raw_pdf, with_chrome=True)

    print("Extracting heading page numbers…")
    heading_pages = extract_heading_pages(raw_pdf)
    html_with_toc = apply_toc_pages(html_doc, heading_pages)

    print("Rendering PDF (pass 2 with TOC pages)…")
    render_pdf_raw(html_with_toc, raw_pdf, with_chrome=True)

    # Embed metadata via pypdf without overlay stamps (Playwright owns H/F)
    from pypdf import PdfReader, PdfWriter

    reader = PdfReader(str(raw_pdf))
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    writer.add_metadata(
        {
            "/Title": f"{brand.DOC_TITLE} — Product Documentation",
            "/Author": brand.AUTHOR,
            "/Subject": "EliteFlow Internship Project — Enterprise Product Documentation",
            "/Creator": "EliteFlow Documentation Generator",
        }
    )
    meta_pdf = paths.ASSETS / "manual-meta.pdf"
    with meta_pdf.open("wb") as f:
        writer.write(f)

    return write_pdf(meta_pdf, paths.OUT_PDF)


def set_run_font(run, name="Calibri", size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def add_toc_field(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-2" \\h \\z \\u '
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr)
    run._r.append(fld_char_separate)
    placeholder = paragraph.add_run("Right-click → Update Field to refresh TOC in Word.")
    set_run_font(placeholder, size=10, color=(75, 85, 99))
    run2 = paragraph.add_run()
    run2._r.append(fld_char_end)


def setup_docx_sections(doc: Document):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    r0 = hp.add_run("EliteFlow")
    set_run_font(r0, size=9, bold=True, color=(11, 58, 110))
    r1 = hp.add_run("  ·  Product Documentation  ·  Internship Project")
    set_run_font(r1, size=9, color=(75, 85, 99))

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    left = fp.add_run("Confidential · EliteFlow   ·   ")
    set_run_font(left, size=8, color=(75, 85, 99))
    add_page_number(fp)
    right = fp.add_run(f"   ·   v{brand.DOC_VERSION}")
    set_run_font(right, size=8, color=(75, 85, 99))


def add_docx_cover(doc: Document):
    p = doc.add_paragraph()
    r = p.add_run("ELITEFLOW")
    set_run_font(r, size=14, bold=True, color=(124, 58, 237))

    title = doc.add_paragraph()
    tr = title.add_run(brand.DOC_TITLE)
    set_run_font(tr, size=26, bold=True, color=(11, 58, 110))

    sub = doc.add_paragraph()
    sr = sub.add_run("Enterprise Product Documentation")
    set_run_font(sr, size=14, color=(21, 101, 192))

    pill = doc.add_paragraph()
    pr = pill.add_run("Internship Project · Software Product Manual")
    set_run_font(pr, size=12, bold=True, color=(17, 24, 39))

    doc.add_paragraph()
    for label, value in [
        ("Prepared By", brand.AUTHOR),
        ("Version", brand.DOC_VERSION),
        ("Date", brand.today_str()),
        ("Classification", "Product Documentation · Internship Submission"),
    ]:
        lp = doc.add_paragraph()
        lr = lp.add_run(f"{label}: ")
        set_run_font(lr, size=11, bold=True, color=(11, 58, 110))
        vr = lp.add_run(value)
        set_run_font(vr, size=11, color=(17, 24, 39))

    doc.add_page_break()


def html_to_docx(html_doc: str) -> None:
    soup = BeautifulSoup(html_doc, "lxml")
    doc = Document()
    setup_docx_sections(doc)
    add_docx_cover(doc)

    h = doc.add_heading("Table of Contents", level=1)
    for run in h.runs:
        set_run_font(run, size=16, bold=True, color=(11, 58, 110))
    toc_p = doc.add_paragraph()
    add_toc_field(toc_p)
    doc.add_page_break()

    body = soup.body
    if not body:
        doc.save(str(paths.OUT_DOCX))
        return

    for el in body.find_all(["h1", "h2", "h3", "p", "ul", "ol", "table", "figure", "pre"], recursive=True):
        parent_names = {p.name for p in el.parents}
        if el.name in ("p", "ul", "ol") and (
            "td" in parent_names
            or "th" in parent_names
            or "figure" in parent_names
            or "li" in parent_names
            or "shot-meta" in (el.parent.get("class") or [] if el.parent else [])
        ):
            # Allow shot-meta paragraphs
            if el.name == "p" and el.parent and "shot-meta" in (el.parent.get("class") or []):
                pass
            elif "shot-meta" not in parent_names:
                if el.name in ("p", "ul", "ol") and (
                    "td" in parent_names or "th" in parent_names or "figure" in parent_names or "li" in parent_names
                ):
                    continue

        if el.name == "h1":
            text = el.get_text(strip=True)
            heading = doc.add_heading(text, level=1)
            for run in heading.runs:
                set_run_font(run, size=16, bold=True, color=(11, 58, 110))
        elif el.name == "h2":
            heading = doc.add_heading(el.get_text(strip=True), level=2)
            for run in heading.runs:
                set_run_font(run, size=13, bold=True, color=(124, 58, 237))
        elif el.name == "h3":
            heading = doc.add_heading(el.get_text(strip=True), level=3)
            for run in heading.runs:
                set_run_font(run, size=11, bold=True, color=(17, 24, 39))
        elif el.name == "p":
            text = el.get_text(" ", strip=True)
            if not text:
                continue
            para = doc.add_paragraph(text)
            for run in para.runs:
                set_run_font(run, size=10, color=(17, 24, 39))
            para.paragraph_format.space_after = Pt(6)
        elif el.name in ("ul", "ol"):
            for li in el.find_all("li", recursive=False):
                para = doc.add_paragraph(li.get_text(" ", strip=True), style="List Bullet")
                for run in para.runs:
                    set_run_font(run, size=10, color=(17, 24, 39))
        elif el.name == "table":
            rows = el.find_all("tr")
            if not rows:
                continue
            cols = max(len(r.find_all(["td", "th"])) for r in rows)
            if cols == 0:
                continue
            table = doc.add_table(rows=len(rows), cols=cols)
            table.style = "Table Grid"
            for ri, row in enumerate(rows):
                cells = row.find_all(["td", "th"])
                for ci in range(cols):
                    cell_text = cells[ci].get_text(" ", strip=True) if ci < len(cells) else ""
                    cell = table.cell(ri, ci)
                    cell.text = cell_text
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            set_run_font(
                                run,
                                size=9,
                                bold=(cells[ci].name == "th") if ci < len(cells) else False,
                                color=(11, 58, 110) if ri == 0 else (17, 24, 39),
                            )
            doc.add_paragraph()
        elif el.name == "figure":
            img = el.find("img")
            cap = el.find("figcaption")
            img_path = None
            if img is not None:
                data_file = img.get("data-file")
                if data_file:
                    candidate = paths.SCREENSHOTS / data_file
                    if candidate.is_file():
                        img_path = candidate
            if img_path and img_path.is_file():
                try:
                    is_portrait = "portrait" in (img.get("class") or [])
                    width = Inches(3.6) if is_portrait else Inches(5.6)
                    doc.add_picture(str(img_path), width=width)
                except Exception as exc:  # noqa: BLE001
                    print(f"Skip image {img_path}: {exc}")
            if cap:
                cp = doc.add_paragraph(cap.get_text(" ", strip=True))
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in cp.runs:
                    set_run_font(run, size=9, bold=True, color=(75, 85, 99))
        elif el.name == "pre":
            para = doc.add_paragraph(el.get_text())
            for run in para.runs:
                set_run_font(run, name="Consolas", size=8, color=(15, 23, 42))

    doc.save(str(paths.OUT_DOCX))
    print(f"DOCX written: {paths.OUT_DOCX}")


def main() -> None:
    print("Building EliteFlow Product Documentation…")
    html_content.figure_counter = 0
    html_content.table_counter = 0
    html_content._PLANNED_TABLES.clear()

    html_doc = build_html()
    paths.ASSETS.mkdir(parents=True, exist_ok=True)
    paths.HTML_PATH.write_text(html_doc, encoding="utf-8")
    print(f"HTML written: {paths.HTML_PATH} ({len(html_doc):,} chars)")

    generate_pdf(html_doc)
    print("Building DOCX…")
    html_content.figure_counter = 0
    html_content.table_counter = 0
    html_to_docx(html_doc)
    print("Done.")


if __name__ == "__main__":
    main()
