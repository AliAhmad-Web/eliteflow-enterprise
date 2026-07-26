#!/usr/bin/env python3
"""
PHASE FP-2 — Generate professional PDF + DOCX from the existing EliteFlow proposal.
Reads docs/FINAL_PROJECT_PROPOSAL_ELITEFLOW.md without modifying it.
"""

from __future__ import annotations

import html
import re
import sys
from datetime import date
from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from markdown import markdown
from playwright.sync_api import sync_playwright

ROOT = Path(__file__).resolve().parents[2]
MD_PATH = ROOT / "docs" / "FINAL_PROJECT_PROPOSAL_ELITEFLOW.md"
OUT_PDF = ROOT / "docs" / "FINAL_PROJECT_PROPOSAL_ELITEFLOW.pdf"
OUT_DOCX = ROOT / "docs" / "FINAL_PROJECT_PROPOSAL_ELITEFLOW.docx"
ASSETS = ROOT / "docs" / "proposal-assets"
HTML_PATH = ASSETS / "proposal.html"

BRAND_BLUE = "#0B3A6E"
BRAND_BLUE_MID = "#1565C0"
BRAND_BLUE_LIGHT = "#E8F1FB"
BRAND_BLACK = "#111827"
BRAND_GRAY = "#4B5563"
BRAND_BORDER = "#D1D5DB"

# Display heading remaps for requested section titles (body text unchanged)
HEADING_REMAP = {
    "1. Executive Summary": "1. Executive Summary",
    "2. Project Vision": "2. Project Vision",
    "3. Problem Statement": "3. Problem Statement",
    "4. Target Audience": "4. Target Audience",
    "5. Business Goals": "5. Business Goals",
    "6. Technical Architecture": "6. Current Architecture",
    "7. Current Completed Features (Phase 1)": "7. Completed Modules",
    "8. Future Roadmap & Module List": "8. Future Roadmap",
    "9. Future AI Vision — AI as the Center of EliteFlow": "9. AI Vision",
    "10. Technology Stack": "10. Technology Stack",
    "11. Development Phases & Timeline": "11. Development Phases",
    "12. Deliverables (This Proposal Package)": "12. Deliverables",
    "13. Risk Analysis": "13. Risk Analysis",
    "14. Scalability Plan": "14. Scalability Strategy",
    "15. Security Plan": "15. Security Strategy",
    "16. Commercial Vision": "16. Commercial Vision",
    "17. Investor-Friendly Summary": "17. Investor Summary",
    "18. Feature List (Consolidated View)": "18. Feature List",
    "19. Closing Statement": "19. Conclusion",
}

SUBHEADING_REMAP = {
    "6.3 Production Architecture (Deployment Strategy)": "6.3 Deployment Architecture",
}


def today_str() -> str:
    return date.today().strftime("%B %d, %Y")


def svg_architecture() -> str:
    return f"""
<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 860 420" width="860" height="420">
  <defs>
    <linearGradient id="g1" x1="0" y1="0" x2="1" y2="1">
      <stop offset="0%" stop-color="#0B3A6E"/><stop offset="100%" stop-color="#1565C0"/>
    </linearGradient>
    <filter id="s" x="-10%" y="-10%" width="120%" height="130%">
      <feDropShadow dx="0" dy="2" stdDeviation="2" flood-opacity="0.15"/>
    </filter>
  </defs>
  <rect width="860" height="420" fill="#F8FAFC"/>
  <text x="30" y="36" font-family="Segoe UI, Arial" font-size="16" font-weight="700" fill="{BRAND_BLUE}">EliteFlow System Architecture</text>
  <g filter="url(#s)">
    <rect x="30" y="60" width="170" height="90" rx="10" fill="url(#g1)"/>
    <text x="115" y="95" text-anchor="middle" fill="#fff" font-family="Segoe UI, Arial" font-size="14" font-weight="700">Web</text>
    <text x="115" y="118" text-anchor="middle" fill="#E8F1FB" font-family="Segoe UI, Arial" font-size="11">Next.js 16 · Vercel</text>
  </g>
  <g filter="url(#s)">
    <rect x="240" y="60" width="170" height="90" rx="10" fill="url(#g1)"/>
    <text x="325" y="95" text-anchor="middle" fill="#fff" font-family="Segoe UI, Arial" font-size="14" font-weight="700">Desktop</text>
    <text x="325" y="118" text-anchor="middle" fill="#E8F1FB" font-family="Segoe UI, Arial" font-size="11">Electron · Shell</text>
  </g>
  <g filter="url(#s)">
    <rect x="450" y="60" width="170" height="90" rx="10" fill="url(#g1)"/>
    <text x="535" y="95" text-anchor="middle" fill="#fff" font-family="Segoe UI, Arial" font-size="14" font-weight="700">Mobile</text>
    <text x="535" y="118" text-anchor="middle" fill="#E8F1FB" font-family="Segoe UI, Arial" font-size="11">Expo 57 · EAS</text>
  </g>
  <g filter="url(#s)">
    <rect x="660" y="60" width="170" height="90" rx="10" fill="url(#g1)"/>
    <text x="745" y="95" text-anchor="middle" fill="#fff" font-family="Segoe UI, Arial" font-size="14" font-weight="700">Extension</text>
    <text x="745" y="118" text-anchor="middle" fill="#E8F1FB" font-family="Segoe UI, Arial" font-size="11">Chrome MV3</text>
  </g>
  <line x1="115" y1="150" x2="115" y2="190" stroke="{BRAND_BLUE}" stroke-width="2"/>
  <line x1="325" y1="150" x2="325" y2="190" stroke="{BRAND_BLUE}" stroke-width="2"/>
  <line x1="535" y1="150" x2="535" y2="190" stroke="{BRAND_BLUE}" stroke-width="2"/>
  <line x1="745" y1="150" x2="745" y2="190" stroke="{BRAND_BLUE}" stroke-width="2"/>
  <line x1="115" y1="190" x2="745" y2="190" stroke="{BRAND_BLUE}" stroke-width="2"/>
  <line x1="430" y1="190" x2="430" y2="220" stroke="{BRAND_BLUE}" stroke-width="2"/>
  <g filter="url(#s)">
    <rect x="280" y="220" width="300" height="80" rx="10" fill="#111827"/>
    <text x="430" y="252" text-anchor="middle" fill="#fff" font-family="Segoe UI, Arial" font-size="15" font-weight="700">API · Express 5 · /api/v1</text>
    <text x="430" y="275" text-anchor="middle" fill="#93C5FD" font-family="Segoe UI, Arial" font-size="12">Railway · JWT + HTTPS</text>
  </g>
  <line x1="330" y1="300" x2="200" y2="340" stroke="{BRAND_BLUE}" stroke-width="2"/>
  <line x1="430" y1="300" x2="430" y2="340" stroke="{BRAND_BLUE}" stroke-width="2"/>
  <line x1="530" y1="300" x2="660" y2="340" stroke="{BRAND_BLUE}" stroke-width="2"/>
  <g filter="url(#s)">
    <rect x="90" y="340" width="220" height="60" rx="8" fill="#E8F1FB" stroke="{BRAND_BLUE_MID}"/>
    <text x="200" y="365" text-anchor="middle" fill="{BRAND_BLUE}" font-family="Segoe UI, Arial" font-size="13" font-weight="700">PostgreSQL · Prisma</text>
    <text x="200" y="385" text-anchor="middle" fill="{BRAND_GRAY}" font-family="Segoe UI, Arial" font-size="11">System of record</text>
  </g>
  <g filter="url(#s)">
    <rect x="320" y="340" width="220" height="60" rx="8" fill="#E8F1FB" stroke="{BRAND_BLUE_MID}"/>
    <text x="430" y="365" text-anchor="middle" fill="{BRAND_BLUE}" font-family="Segoe UI, Arial" font-size="13" font-weight="700">Supabase · Resend</text>
    <text x="430" y="385" text-anchor="middle" fill="{BRAND_GRAY}" font-family="Segoe UI, Arial" font-size="11">OAuth · Storage · Email</text>
  </g>
  <g filter="url(#s)">
    <rect x="550" y="340" width="220" height="60" rx="8" fill="#E8F1FB" stroke="{BRAND_BLUE_MID}"/>
    <text x="660" y="365" text-anchor="middle" fill="{BRAND_BLUE}" font-family="Segoe UI, Arial" font-size="13" font-weight="700">AI Providers</text>
    <text x="660" y="385" text-anchor="middle" fill="{BRAND_GRAY}" font-family="Segoe UI, Arial" font-size="11">Gemini · OpenAI · Mock</text>
  </g>
</svg>
"""


def svg_deployment() -> str:
    return f"""
<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 860 300" width="860" height="300">
  <rect width="860" height="300" fill="#F8FAFC"/>
  <text x="30" y="36" font-family="Segoe UI, Arial" font-size="16" font-weight="700" fill="{BRAND_BLUE}">Deployment Flow</text>
  <rect x="30" y="70" width="140" height="70" rx="8" fill="{BRAND_BLUE}"/>
  <text x="100" y="100" text-anchor="middle" fill="#fff" font-family="Segoe UI, Arial" font-size="13" font-weight="700">Git Monorepo</text>
  <text x="100" y="120" text-anchor="middle" fill="#BFDBFE" font-family="Segoe UI, Arial" font-size="11">apps + packages</text>
  <polygon points="185,105 205,95 205,115" fill="{BRAND_BLUE_MID}"/>
  <line x1="170" y1="105" x2="185" y2="105" stroke="{BRAND_BLUE_MID}" stroke-width="2"/>
  <rect x="220" y="55" width="160" height="50" rx="8" fill="#111827"/>
  <text x="300" y="85" text-anchor="middle" fill="#fff" font-family="Segoe UI, Arial" font-size="12" font-weight="700">Vercel → Web</text>
  <rect x="220" y="120" width="160" height="50" rx="8" fill="#111827"/>
  <text x="300" y="150" text-anchor="middle" fill="#fff" font-family="Segoe UI, Arial" font-size="12" font-weight="700">Railway → API</text>
  <polygon points="395,80 415,70 415,90" fill="{BRAND_BLUE_MID}"/>
  <polygon points="395,145 415,135 415,155" fill="{BRAND_BLUE_MID}"/>
  <rect x="430" y="55" width="180" height="50" rx="8" fill="{BRAND_BLUE_LIGHT}" stroke="{BRAND_BLUE_MID}"/>
  <text x="520" y="85" text-anchor="middle" fill="{BRAND_BLUE}" font-family="Segoe UI, Arial" font-size="12" font-weight="700">CDN / Edge</text>
  <rect x="430" y="120" width="180" height="50" rx="8" fill="{BRAND_BLUE_LIGHT}" stroke="{BRAND_BLUE_MID}"/>
  <text x="520" y="150" text-anchor="middle" fill="{BRAND_BLUE}" font-family="Segoe UI, Arial" font-size="12" font-weight="700">Postgres · Supabase</text>
  <rect x="30" y="210" width="800" height="60" rx="8" fill="#fff" stroke="{BRAND_BORDER}"/>
  <text x="50" y="245" fill="{BRAND_BLACK}" font-family="Segoe UI, Arial" font-size="13" font-weight="600">Clients:</text>
  <text x="130" y="245" fill="{BRAND_GRAY}" font-family="Segoe UI, Arial" font-size="13">Web · Desktop · Android · iOS · Chrome Extension  →  same /api/v1</text>
</svg>
"""


def svg_roadmap() -> str:
    return f"""
<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 860 220" width="860" height="220">
  <rect width="860" height="220" fill="#F8FAFC"/>
  <text x="30" y="36" font-family="Segoe UI, Arial" font-size="16" font-weight="700" fill="{BRAND_BLUE}">Development Phase Roadmap</text>
  <line x1="60" y1="120" x2="800" y2="120" stroke="{BRAND_BORDER}" stroke-width="4"/>
  <g>
    <circle cx="120" cy="120" r="18" fill="#059669"/>
    <text x="120" y="125" text-anchor="middle" fill="#fff" font-family="Segoe UI, Arial" font-size="12" font-weight="700">1</text>
    <text x="120" y="70" text-anchor="middle" fill="#059669" font-family="Segoe UI, Arial" font-size="13" font-weight="700">Phase 1</text>
    <text x="120" y="90" text-anchor="middle" fill="{BRAND_GRAY}" font-family="Segoe UI, Arial" font-size="11">Completed</text>
    <text x="120" y="165" text-anchor="middle" fill="{BRAND_BLACK}" font-family="Segoe UI, Arial" font-size="11">Foundation EMS</text>
  </g>
  <g>
    <circle cx="320" cy="120" r="18" fill="{BRAND_BLUE}"/>
    <text x="320" y="125" text-anchor="middle" fill="#fff" font-family="Segoe UI, Arial" font-size="12" font-weight="700">2</text>
    <text x="320" y="70" text-anchor="middle" fill="{BRAND_BLUE}" font-family="Segoe UI, Arial" font-size="13" font-weight="700">Phase 2</text>
    <text x="320" y="90" text-anchor="middle" fill="{BRAND_GRAY}" font-family="Segoe UI, Arial" font-size="11">Next</text>
    <text x="320" y="165" text-anchor="middle" fill="{BRAND_BLACK}" font-family="Segoe UI, Arial" font-size="11">Enterprise Features</text>
  </g>
  <g>
    <circle cx="520" cy="120" r="18" fill="{BRAND_BLUE_MID}"/>
    <text x="520" y="125" text-anchor="middle" fill="#fff" font-family="Segoe UI, Arial" font-size="12" font-weight="700">3</text>
    <text x="520" y="70" text-anchor="middle" fill="{BRAND_BLUE_MID}" font-family="Segoe UI, Arial" font-size="13" font-weight="700">Phase 3</text>
    <text x="520" y="90" text-anchor="middle" fill="{BRAND_GRAY}" font-family="Segoe UI, Arial" font-size="11">AI Wave</text>
    <text x="520" y="165" text-anchor="middle" fill="{BRAND_BLACK}" font-family="Segoe UI, Arial" font-size="11">AI Expansion</text>
  </g>
  <g>
    <circle cx="720" cy="120" r="18" fill="#111827"/>
    <text x="720" y="125" text-anchor="middle" fill="#fff" font-family="Segoe UI, Arial" font-size="12" font-weight="700">4</text>
    <text x="720" y="70" text-anchor="middle" fill="#111827" font-family="Segoe UI, Arial" font-size="13" font-weight="700">Phase 4</text>
    <text x="720" y="90" text-anchor="middle" fill="{BRAND_GRAY}" font-family="Segoe UI, Arial" font-size="11">Commercial</text>
    <text x="720" y="165" text-anchor="middle" fill="{BRAND_BLACK}" font-family="Segoe UI, Arial" font-size="11">Commercial SaaS</text>
  </g>
</svg>
"""


def svg_tech_stack() -> str:
    return f"""
<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 860 280" width="860" height="280">
  <rect width="860" height="280" fill="#F8FAFC"/>
  <text x="30" y="36" font-family="Segoe UI, Arial" font-size="16" font-weight="700" fill="{BRAND_BLUE}">Technology Stack Overview</text>
  <rect x="30" y="60" width="190" height="190" rx="10" fill="#fff" stroke="{BRAND_BORDER}"/>
  <rect x="30" y="60" width="190" height="36" rx="10" fill="{BRAND_BLUE}"/>
  <text x="125" y="84" text-anchor="middle" fill="#fff" font-family="Segoe UI, Arial" font-size="13" font-weight="700">Clients</text>
  <text x="125" y="120" text-anchor="middle" fill="{BRAND_BLACK}" font-family="Segoe UI, Arial" font-size="12">Next.js 16</text>
  <text x="125" y="145" text-anchor="middle" fill="{BRAND_BLACK}" font-family="Segoe UI, Arial" font-size="12">Electron</text>
  <text x="125" y="170" text-anchor="middle" fill="{BRAND_BLACK}" font-family="Segoe UI, Arial" font-size="12">Expo 57</text>
  <text x="125" y="195" text-anchor="middle" fill="{BRAND_BLACK}" font-family="Segoe UI, Arial" font-size="12">Chrome MV3</text>
  <text x="125" y="220" text-anchor="middle" fill="{BRAND_GRAY}" font-family="Segoe UI, Arial" font-size="11">React 19 · RN</text>
  <rect x="240" y="60" width="190" height="190" rx="10" fill="#fff" stroke="{BRAND_BORDER}"/>
  <rect x="240" y="60" width="190" height="36" rx="10" fill="{BRAND_BLUE}"/>
  <text x="335" y="84" text-anchor="middle" fill="#fff" font-family="Segoe UI, Arial" font-size="13" font-weight="700">API</text>
  <text x="335" y="120" text-anchor="middle" fill="{BRAND_BLACK}" font-family="Segoe UI, Arial" font-size="12">Express 5</text>
  <text x="335" y="145" text-anchor="middle" fill="{BRAND_BLACK}" font-family="Segoe UI, Arial" font-size="12">TypeScript</text>
  <text x="335" y="170" text-anchor="middle" fill="{BRAND_BLACK}" font-family="Segoe UI, Arial" font-size="12">JWT · Argon2</text>
  <text x="335" y="195" text-anchor="middle" fill="{BRAND_BLACK}" font-family="Segoe UI, Arial" font-size="12">Zod contracts</text>
  <text x="335" y="220" text-anchor="middle" fill="{BRAND_GRAY}" font-family="Segoe UI, Arial" font-size="11">/api/v1</text>
  <rect x="450" y="60" width="190" height="190" rx="10" fill="#fff" stroke="{BRAND_BORDER}"/>
  <rect x="450" y="60" width="190" height="36" rx="10" fill="{BRAND_BLUE}"/>
  <text x="545" y="84" text-anchor="middle" fill="#fff" font-family="Segoe UI, Arial" font-size="13" font-weight="700">Data &amp; AI</text>
  <text x="545" y="120" text-anchor="middle" fill="{BRAND_BLACK}" font-family="Segoe UI, Arial" font-size="12">PostgreSQL</text>
  <text x="545" y="145" text-anchor="middle" fill="{BRAND_BLACK}" font-family="Segoe UI, Arial" font-size="12">Prisma 6</text>
  <text x="545" y="170" text-anchor="middle" fill="{BRAND_BLACK}" font-family="Segoe UI, Arial" font-size="12">Supabase</text>
  <text x="545" y="195" text-anchor="middle" fill="{BRAND_BLACK}" font-family="Segoe UI, Arial" font-size="12">Gemini / OpenAI</text>
  <text x="545" y="220" text-anchor="middle" fill="{BRAND_GRAY}" font-family="Segoe UI, Arial" font-size="11">Resend</text>
  <rect x="660" y="60" width="170" height="190" rx="10" fill="#fff" stroke="{BRAND_BORDER}"/>
  <rect x="660" y="60" width="170" height="36" rx="10" fill="#111827"/>
  <text x="745" y="84" text-anchor="middle" fill="#fff" font-family="Segoe UI, Arial" font-size="13" font-weight="700">Deploy</text>
  <text x="745" y="120" text-anchor="middle" fill="{BRAND_BLACK}" font-family="Segoe UI, Arial" font-size="12">Vercel</text>
  <text x="745" y="145" text-anchor="middle" fill="{BRAND_BLACK}" font-family="Segoe UI, Arial" font-size="12">Railway</text>
  <text x="745" y="170" text-anchor="middle" fill="{BRAND_BLACK}" font-family="Segoe UI, Arial" font-size="12">EAS</text>
  <text x="745" y="195" text-anchor="middle" fill="{BRAND_BLACK}" font-family="Segoe UI, Arial" font-size="12">CDN / Storage</text>
  <text x="745" y="220" text-anchor="middle" fill="{BRAND_GRAY}" font-family="Segoe UI, Arial" font-size="11">npm workspaces</text>
</svg>
"""


def svg_modules() -> str:
    return f"""
<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 860 320" width="860" height="320">
  <rect width="860" height="320" fill="#F8FAFC"/>
  <text x="30" y="36" font-family="Segoe UI, Arial" font-size="16" font-weight="700" fill="{BRAND_BLUE}">Module Hierarchy</text>
  <rect x="300" y="55" width="260" height="50" rx="10" fill="{BRAND_BLUE}"/>
  <text x="430" y="85" text-anchor="middle" fill="#fff" font-family="Segoe UI, Arial" font-size="14" font-weight="700">EliteFlow Core Platform</text>
  <line x1="430" y1="105" x2="430" y2="135" stroke="{BRAND_BLUE}" stroke-width="2"/>
  <line x1="110" y1="135" x2="750" y2="135" stroke="{BRAND_BLUE}" stroke-width="2"/>
  <line x1="110" y1="135" x2="110" y2="155" stroke="{BRAND_BLUE}" stroke-width="2"/>
  <line x1="290" y1="135" x2="290" y2="155" stroke="{BRAND_BLUE}" stroke-width="2"/>
  <line x1="470" y1="135" x2="470" y2="155" stroke="{BRAND_BLUE}" stroke-width="2"/>
  <line x1="650" y1="135" x2="650" y2="155" stroke="{BRAND_BLUE}" stroke-width="2"/>
  <line x1="750" y1="135" x2="750" y2="155" stroke="{BRAND_BLUE}" stroke-width="2"/>
  <rect x="40" y="155" width="140" height="44" rx="8" fill="#059669"/><text x="110" y="182" text-anchor="middle" fill="#fff" font-size="12" font-family="Segoe UI, Arial" font-weight="700">Phase 1 Live</text>
  <rect x="220" y="155" width="140" height="44" rx="8" fill="{BRAND_BLUE}"/><text x="290" y="182" text-anchor="middle" fill="#fff" font-size="12" font-family="Segoe UI, Arial" font-weight="700">Growth Ops</text>
  <rect x="400" y="155" width="140" height="44" rx="8" fill="{BRAND_BLUE_MID}"/><text x="470" y="182" text-anchor="middle" fill="#fff" font-size="12" font-family="Segoe UI, Arial" font-weight="700">People / ERP</text>
  <rect x="580" y="155" width="140" height="44" rx="8" fill="#111827"/><text x="650" y="182" text-anchor="middle" fill="#fff" font-size="12" font-family="Segoe UI, Arial" font-weight="700">AI Platform</text>
  <rect x="700" y="155" width="120" height="44" rx="8" fill="#374151"/><text x="760" y="182" text-anchor="middle" fill="#fff" font-size="11" font-family="Segoe UI, Arial" font-weight="700">Ecosystem</text>
  <text x="110" y="230" text-anchor="middle" fill="{BRAND_GRAY}" font-size="11" font-family="Segoe UI, Arial">Clients · Projects</text>
  <text x="110" y="248" text-anchor="middle" fill="{BRAND_GRAY}" font-size="11" font-family="Segoe UI, Arial">Tasks · AI · Comms</text>
  <text x="290" y="230" text-anchor="middle" fill="{BRAND_GRAY}" font-size="11" font-family="Segoe UI, Arial">CRM · Pipeline</text>
  <text x="290" y="248" text-anchor="middle" fill="{BRAND_GRAY}" font-size="11" font-family="Segoe UI, Arial">Support · Marketing</text>
  <text x="470" y="230" text-anchor="middle" fill="{BRAND_GRAY}" font-size="11" font-family="Segoe UI, Arial">HRM · Payroll</text>
  <text x="470" y="248" text-anchor="middle" fill="{BRAND_GRAY}" font-size="11" font-family="Segoe UI, Arial">Inventory · Finance</text>
  <text x="650" y="230" text-anchor="middle" fill="{BRAND_GRAY}" font-size="11" font-family="Segoe UI, Arial">Automation</text>
  <text x="650" y="248" text-anchor="middle" fill="{BRAND_GRAY}" font-size="11" font-family="Segoe UI, Arial">Insights · OCR</text>
  <text x="760" y="230" text-anchor="middle" fill="{BRAND_GRAY}" font-size="11" font-family="Segoe UI, Arial">API · Plugins</text>
  <text x="760" y="248" text-anchor="middle" fill="{BRAND_GRAY}" font-size="11" font-family="Segoe UI, Arial">Integrations</text>
  <text x="430" y="295" text-anchor="middle" fill="{BRAND_BLUE}" font-size="12" font-family="Segoe UI, Arial" font-weight="600">Extend EliteFlow — never replace the foundation</text>
</svg>
"""


def preprocess_markdown(md: str) -> str:
    """Apply display heading remaps only; preserve all body content."""
    for old, new in HEADING_REMAP.items():
        md = md.replace(f"# {old}", f"# {new}", 1)
    for old, new in SUBHEADING_REMAP.items():
        md = md.replace(f"## {old}", f"## {new}", 1)
    return md


def md_to_body_html(md: str) -> str:
    body = markdown(
        md,
        extensions=["tables", "fenced_code", "sane_lists", "toc", "nl2br"],
        extension_configs={"toc": {"permalink": False, "toc_depth": "2-3"}},
    )
    soup = BeautifulSoup(body, "lxml")

    # Remap h1 ids for TOC anchors
    for h in soup.find_all(["h1", "h2", "h3"]):
        text = h.get_text(strip=True)
        slug = re.sub(r"[^a-z0-9]+", "-", text.lower()).strip("-")
        h["id"] = slug

    # Insert diagrams after key sections
    inserts = [
        ("6-current-architecture", "Architecture Diagram", svg_architecture(), "after_h1"),
        ("6-3-deployment-architecture", "Deployment Architecture Diagram", svg_deployment(), "after_h2"),
        ("8-future-roadmap", "Module Hierarchy", svg_modules(), "after_h1"),
        ("10-technology-stack", "Technology Stack Diagram", svg_tech_stack(), "after_h1"),
        ("11-development-phases", "Phase Roadmap", svg_roadmap(), "after_h1"),
    ]

    for anchor, title, svg, _ in inserts:
        target = soup.find(id=anchor)
        if not target:
            continue
        figure = soup.new_tag("figure", attrs={"class": "diagram"})
        cap = soup.new_tag("figcaption")
        cap.string = title
        figure.append(BeautifulSoup(svg, "xml"))
        figure.append(cap)
        target.insert_after(figure)

    # Extract only body children
    body_el = soup.body
    return "".join(str(c) for c in body_el.children) if body_el else str(soup)


def build_toc_entries(md: str) -> list[tuple[str, str, int]]:
    entries = []
    for line in md.splitlines():
        if line.startswith("# "):
            title = line[2:].strip()
            slug = re.sub(r"[^a-z0-9]+", "-", title.lower()).strip("-")
            entries.append((title, slug, 1))
        elif line.startswith("## ") and not line.startswith("### "):
            title = line[3:].strip()
            # Include key subsections only
            if title.startswith("6.") or title.startswith("10.") or "Deployment" in title:
                slug = re.sub(r"[^a-z0-9]+", "-", title.lower()).strip("-")
                entries.append((title, slug, 2))
    return entries


def css() -> str:
    return f"""
@page {{
  size: A4;
  margin: 0;
}}
* {{ box-sizing: border-box; }}
html, body {{
  margin: 0;
  padding: 0;
  font-family: "Segoe UI", "Calibri", Arial, sans-serif;
  color: {BRAND_BLACK};
  font-size: 10.5pt;
  line-height: 1.55;
  background: #fff;
}}
.cover {{
  page-break-after: always;
  min-height: 297mm;
  width: 210mm;
  margin: 0;
  padding: 0;
  background: linear-gradient(165deg, #061F3B 0%, {BRAND_BLUE} 45%, #1E4E8C 100%);
  color: #fff;
  position: relative;
  overflow: hidden;
}}
.cover::before {{
  content: "";
  position: absolute;
  right: -80px;
  top: -80px;
  width: 320px;
  height: 320px;
  border-radius: 50%;
  background: rgba(255,255,255,0.06);
}}
.cover::after {{
  content: "";
  position: absolute;
  left: -60px;
  bottom: 80px;
  width: 220px;
  height: 220px;
  border-radius: 50%;
  background: rgba(21,101,192,0.35);
}}
.cover-inner {{
  position: relative;
  z-index: 1;
  padding: 48mm 22mm 30mm 22mm;
  min-height: 297mm;
  display: flex;
  flex-direction: column;
}}
.brand-mark {{
  display: inline-flex;
  align-items: center;
  gap: 12px;
  margin-bottom: 48px;
}}
.brand-icon {{
  width: 48px;
  height: 48px;
  border-radius: 12px;
  background: #fff;
  color: {BRAND_BLUE};
  display: flex;
  align-items: center;
  justify-content: center;
  font-weight: 800;
  font-size: 18px;
}}
.brand-name {{
  font-size: 14pt;
  letter-spacing: 0.12em;
  text-transform: uppercase;
  opacity: 0.9;
}}
.cover h1 {{
  font-size: 34pt;
  line-height: 1.15;
  margin: 0 0 12px 0;
  font-weight: 700;
  max-width: 85%;
  color: #fff;
  border: none;
  padding: 0;
}}
.cover .subtitle {{
  font-size: 16pt;
  font-weight: 400;
  opacity: 0.95;
  margin: 0 0 8px 0;
}}
.cover .pill {{
  display: inline-block;
  margin-top: 18px;
  padding: 8px 16px;
  border: 1px solid rgba(255,255,255,0.35);
  border-radius: 999px;
  font-size: 11pt;
  letter-spacing: 0.04em;
}}
.cover-meta {{
  margin-top: auto;
  display: grid;
  grid-template-columns: 1fr 1fr;
  gap: 18px 28px;
  padding-top: 40px;
  border-top: 1px solid rgba(255,255,255,0.25);
}}
.cover-meta .label {{
  font-size: 9pt;
  text-transform: uppercase;
  letter-spacing: 0.08em;
  opacity: 0.7;
  margin-bottom: 4px;
}}
.cover-meta .value {{
  font-size: 13pt;
  font-weight: 600;
}}
.toc-page {{
  page-break-after: always;
  padding: 18mm 16mm 22mm 16mm;
}}
.toc-page h1 {{
  color: {BRAND_BLUE};
  border-bottom: 2px solid {BRAND_BLUE};
  padding-bottom: 8px;
  margin-top: 0;
}}
.content {{
  padding: 18mm 16mm 22mm 16mm;
}}
.toc-list {{
  list-style: none;
  padding: 0;
  margin: 18px 0 0 0;
}}
.toc-list li {{
  display: flex;
  align-items: baseline;
  gap: 8px;
  margin: 8px 0;
  font-size: 11pt;
}}
.toc-list li.level-2 {{
  padding-left: 18px;
  font-size: 10pt;
  color: {BRAND_GRAY};
}}
.toc-list a {{
  color: {BRAND_BLACK};
  text-decoration: none;
  flex: 0 1 auto;
}}
.toc-list .dots {{
  flex: 1 1 auto;
  border-bottom: 1px dotted {BRAND_BORDER};
  height: 0.9em;
  min-width: 20px;
}}
.toc-list .page-ref {{
  color: {BRAND_BLUE};
  font-variant-numeric: tabular-nums;
}}
.content h1 {{
  color: {BRAND_BLUE};
  font-size: 18pt;
  margin: 22pt 0 10pt 0;
  padding-bottom: 6px;
  border-bottom: 2px solid {BRAND_BLUE_LIGHT};
  page-break-after: avoid;
  page-break-before: always;
}}
.content > h1:first-of-type,
.content > .meta-strip + h1 {{
  page-break-before: auto;
  margin-top: 0;
}}
.content h2 {{
  color: {BRAND_BLUE_MID};
  font-size: 13.5pt;
  margin: 16pt 0 8pt 0;
  page-break-after: avoid;
}}
.content h3 {{
  color: {BRAND_BLACK};
  font-size: 11.5pt;
  margin: 12pt 0 6pt 0;
  page-break-after: avoid;
}}
.content p {{
  margin: 0 0 9pt 0;
  orphans: 3;
  widows: 3;
}}
.content ul, .content ol {{
  margin: 0 0 10pt 0;
  padding-left: 22px;
}}
.content li {{
  margin: 3pt 0;
}}
.content strong {{
  color: {BRAND_BLACK};
}}
.content table {{
  width: 100%;
  border-collapse: collapse;
  margin: 10pt 0 14pt 0;
  font-size: 9.5pt;
  page-break-inside: avoid;
}}
.content th {{
  background: {BRAND_BLUE};
  color: #fff;
  text-align: left;
  padding: 8px 10px;
  font-weight: 600;
}}
.content td {{
  border: 1px solid {BRAND_BORDER};
  padding: 7px 10px;
  vertical-align: top;
}}
.content tr:nth-child(even) td {{
  background: #F8FAFC;
}}
.content pre {{
  background: #0F172A;
  color: #E2E8F0;
  padding: 12px 14px;
  border-radius: 8px;
  font-family: Consolas, "Courier New", monospace;
  font-size: 8pt;
  line-height: 1.4;
  overflow-x: auto;
  white-space: pre-wrap;
  page-break-inside: avoid;
}}
.content code {{
  font-family: Consolas, "Courier New", monospace;
  font-size: 9pt;
  background: {BRAND_BLUE_LIGHT};
  color: {BRAND_BLUE};
  padding: 1px 5px;
  border-radius: 4px;
}}
.content pre code {{
  background: transparent;
  color: inherit;
  padding: 0;
}}
.content blockquote {{
  margin: 10pt 0;
  padding: 8pt 14pt;
  border-left: 4px solid {BRAND_BLUE};
  background: {BRAND_BLUE_LIGHT};
  color: {BRAND_GRAY};
}}
.content hr {{
  border: none;
  border-top: 1px solid {BRAND_BORDER};
  margin: 16pt 0;
}}
figure.diagram {{
  margin: 14pt 0 18pt 0;
  page-break-inside: avoid;
  background: #fff;
  border: 1px solid {BRAND_BORDER};
  border-radius: 10px;
  padding: 10px;
}}
figure.diagram svg {{
  width: 100%;
  height: auto;
  display: block;
}}
figure.diagram figcaption {{
  text-align: center;
  font-size: 9pt;
  color: {BRAND_GRAY};
  margin-top: 6px;
  font-weight: 600;
}}
.meta-strip {{
  background: {BRAND_BLUE_LIGHT};
  border: 1px solid {BRAND_BORDER};
  border-radius: 8px;
  padding: 10px 14px;
  margin: 0 0 16px 0;
  font-size: 9.5pt;
  color: {BRAND_GRAY};
}}
"""


def build_html(md_raw: str) -> str:
    md = preprocess_markdown(md_raw)
    # Drop the duplicate top H1/H2 title block from MD to avoid repeating cover content awkwardly —
    # keep product meta paragraph content by converting first lines carefully.
    # We keep ALL section content; only skip the very first two heading lines if they match title.
    lines = md.splitlines()
    start = 0
    if lines and lines[0].startswith("# EliteFlow"):
        start = 1
        while start < len(lines) and (
            lines[start].startswith("## ")
            or lines[start].strip() == ""
            or lines[start].startswith("**Product:**")
            or lines[start].startswith("**Version")
            or lines[start].startswith("**Document")
            or lines[start].startswith("**Audience")
            or lines[start].startswith("**Principle")
            or lines[start].strip() == "---"
        ):
            start += 1
    # Preserve meta as intro strip text (content not lost)
    meta_bits = []
    for line in md.splitlines()[:20]:
        if line.startswith("**Product:**"):
            meta_bits.append(line.strip("*").replace("Product:", "").strip())
        if line.startswith("**Principle:**"):
            meta_bits.append(line.replace("**Principle:**", "Principle:").replace("*", "").strip())

    body_md = "\n".join(lines[start:])
    body_html = md_to_body_html(body_md)
    toc = build_toc_entries(body_md)

    toc_html = ['<ol class="toc-list">']
    for title, slug, level in toc:
        cls = f' class="level-{level}"' if level > 1 else ""
        toc_html.append(
            f'<li{cls}><a href="#{html.escape(slug)}">{html.escape(title)}</a>'
            f'<span class="dots"></span><span class="page-ref"></span></li>'
        )
    toc_html.append("</ol>")

    meta_strip = ""
    if meta_bits:
        meta_strip = f'<div class="meta-strip">{" · ".join(html.escape(b) for b in meta_bits)}</div>'

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="utf-8"/>
  <title>EliteFlow Enterprise Platform — Final Project Proposal</title>
  <style>{css()}</style>
</head>
<body>
  <section class="cover">
    <div class="cover-inner">
      <div class="brand-mark">
        <div class="brand-icon">EF</div>
        <div class="brand-name">EliteFlow</div>
      </div>
      <h1>EliteFlow Enterprise Platform</h1>
      <p class="subtitle">Final Project Proposal</p>
      <div class="pill">Internship Project</div>
      <div class="cover-meta">
        <div>
          <div class="label">Prepared By</div>
          <div class="value">Ali Ahmad</div>
        </div>
        <div>
          <div class="label">Version</div>
          <div class="value">1.0</div>
        </div>
        <div>
          <div class="label">Date</div>
          <div class="value">{html.escape(today_str())}</div>
        </div>
        <div>
          <div class="label">Classification</div>
          <div class="value">Mentor · Client · Investor Ready</div>
        </div>
      </div>
    </div>
  </section>

  <section class="toc-page">
    <h1>Table of Contents</h1>
    {''.join(toc_html)}
  </section>

  <main class="content">
    {meta_strip}
    {body_html}
  </main>
</body>
</html>
"""


def extract_heading_pages(pdf_path: Path) -> dict[str, int]:
    """Map section titles to 1-based PDF page numbers using large heading fonts."""
    import fitz

    doc = fitz.open(str(pdf_path))
    mapping: dict[str, int] = {}
    for i, page in enumerate(doc):
        blocks = page.get_text("dict")["blocks"]
        for block in blocks:
            for line in block.get("lines", []):
                spans = line.get("spans", [])
                if not spans:
                    continue
                size = max(s.get("size", 0) for s in spans)
                text = "".join(s.get("text", "") for s in spans).strip()
                if size < 13 or not text:
                    continue
                # Main numbered sections and key TOC titles
                if text.startswith(tuple(f"{n}. " for n in range(1, 20))) or text == "Table of Contents":
                    if size >= 15:
                        mapping.setdefault(text, i + 1)
                # Subsection 6.x / 10.x
                if text.startswith(("6.1 ", "6.2 ", "6.3 ", "10.1 ", "10.2 ")):
                    mapping.setdefault(text, i + 1)
    return mapping


def apply_toc_pages(html_doc: str, heading_pages: dict[str, int]) -> str:
    soup = BeautifulSoup(html_doc, "lxml")
    for li in soup.select(".toc-list li"):
        a = li.find("a")
        ref = li.find(class_="page-ref")
        if not a or ref is None:
            continue
        title = a.get_text(strip=True)
        page_no = heading_pages.get(title)
        if page_no:
            ref.string = str(page_no)
    return str(soup)


def generate_pdf(html_doc: str) -> None:
    from io import BytesIO

    from pypdf import PdfReader, PdfWriter
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas as pdfcanvas

    ASSETS.mkdir(parents=True, exist_ok=True)

    def render_raw(doc_html: str, dest: Path) -> None:
        HTML_PATH.write_text(doc_html, encoding="utf-8")
        with sync_playwright() as p:
            browser = p.chromium.launch()
            page = browser.new_page()
            page.goto(HTML_PATH.as_uri(), wait_until="networkidle")
            page.pdf(
                path=str(dest),
                format="A4",
                print_background=True,
                display_header_footer=False,
                margin={"top": "0", "bottom": "0", "left": "0", "right": "0"},
                prefer_css_page_size=True,
            )
            browser.close()

    raw_pdf = ASSETS / "raw.pdf"
    render_raw(html_doc, raw_pdf)

    # Second pass: accurate TOC page numbers from rendered PDF headings
    heading_pages = extract_heading_pages(raw_pdf)
    html_with_toc = apply_toc_pages(html_doc, heading_pages)
    render_raw(html_with_toc, raw_pdf)

    reader = PdfReader(str(raw_pdf))
    total = len(reader.pages)
    width, height = A4
    stamped = PdfWriter()
    for i, pdf_page in enumerate(reader.pages):
        if i == 0:
            stamped.add_page(pdf_page)
            continue
        packet = BytesIO()
        c = pdfcanvas.Canvas(packet, pagesize=A4)
        c.setFont("Helvetica", 8)
        c.setFillColorRGB(0.043, 0.227, 0.431)
        c.drawString(40, height - 22, "EliteFlow Enterprise Platform")
        c.setFillColorRGB(0.294, 0.333, 0.388)
        c.drawRightString(width - 40, height - 22, "Final Project Proposal")
        c.setStrokeColorRGB(0.82, 0.835, 0.855)
        c.setLineWidth(0.4)
        c.line(40, height - 28, width - 40, height - 28)
        c.line(40, 28, width - 40, 28)
        c.setFillColorRGB(0.294, 0.333, 0.388)
        c.drawString(40, 16, "Confidential · Internship Project")
        c.drawCentredString(width / 2, 16, f"Page {i + 1} of {total}")
        c.drawRightString(width - 40, 16, "v1.0")
        c.save()
        packet.seek(0)
        overlay = PdfReader(packet).pages[0]
        pdf_page.merge_page(overlay)
        stamped.add_page(pdf_page)

    stamped.add_metadata(
        {
            "/Title": "EliteFlow Enterprise Platform — Final Project Proposal",
            "/Author": "Ali Ahmad",
            "/Subject": "Internship Final Project Proposal",
            "/Creator": "EliteFlow Proposal Generator",
        }
    )
    with OUT_PDF.open("wb") as f:
        stamped.write(f)


def set_run_font(run, name="Calibri", size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def add_toc_field(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-2" \\h \\z \\u '
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr)
    run._r.append(fld_char_separate)
    placeholder = paragraph.add_run("Right-click → Update Field to refresh TOC in Word.")
    set_run_font(placeholder, size=10, color=(75, 85, 99))
    run2 = paragraph.add_run()
    run2._r.append(fld_char_end)


def setup_docx_header_footer(doc: Document):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.text = "EliteFlow Enterprise Platform"
    set_run_font(hp.runs[0], size=9, bold=True, color=(11, 58, 110))
    hp.add_run("  |  Final Project Proposal")
    set_run_font(hp.runs[1], size=9, color=(75, 85, 99))

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    left = fp.add_run("Confidential · Internship Project   ·   ")
    set_run_font(left, size=8, color=(75, 85, 99))
    add_page_number(fp)
    right = fp.add_run("   ·   v1.0")
    set_run_font(right, size=8, color=(75, 85, 99))


def add_cover_docx(doc: Document):
    # Brand
    p = doc.add_paragraph()
    r = p.add_run("ELITEFLOW")
    set_run_font(r, size=14, bold=True, color=(11, 58, 110))

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    tr = title.add_run("EliteFlow Enterprise Platform")
    set_run_font(tr, size=28, bold=True, color=(11, 58, 110))

    sub = doc.add_paragraph()
    sr = sub.add_run("Final Project Proposal")
    set_run_font(sr, size=16, color=(21, 101, 192))

    pill = doc.add_paragraph()
    pr = pill.add_run("Internship Project")
    set_run_font(pr, size=12, bold=True, color=(17, 24, 39))

    doc.add_paragraph()
    for label, value in [
        ("Prepared By", "Ali Ahmad"),
        ("Version", "1.0"),
        ("Date", today_str()),
        ("Classification", "Mentor · Client · Investor Ready"),
    ]:
        lp = doc.add_paragraph()
        lr = lp.add_run(f"{label}: ")
        set_run_font(lr, size=11, bold=True, color=(11, 58, 110))
        vr = lp.add_run(value)
        set_run_font(vr, size=11, color=(17, 24, 39))

    doc.add_page_break()


def html_inline_to_runs(paragraph, element, base_size=11):
    if element.name is None:
        text = str(element)
        if text.strip():
            run = paragraph.add_run(text)
            set_run_font(run, size=base_size)
        return
    if element.name in ("strong", "b"):
        run = paragraph.add_run(element.get_text())
        set_run_font(run, size=base_size, bold=True)
        return
    if element.name in ("em", "i"):
        run = paragraph.add_run(element.get_text())
        set_run_font(run, size=base_size)
        run.italic = True
        return
    if element.name == "code":
        run = paragraph.add_run(element.get_text())
        set_run_font(run, name="Consolas", size=9, color=(11, 58, 110))
        return
    if element.name == "br":
        paragraph.add_run().add_break()
        return
    for child in element.children:
        html_inline_to_runs(paragraph, child, base_size)


def generate_docx(md_raw: str) -> None:
    md = preprocess_markdown(md_raw)
    lines = md.splitlines()
    start = 0
    if lines and lines[0].startswith("# EliteFlow"):
        start = 1
        while start < len(lines) and (
            lines[start].startswith("## ")
            or lines[start].strip() == ""
            or lines[start].startswith("**")
            or lines[start].strip() == "---"
        ):
            start += 1
    body_md = "\n".join(lines[start:])
    body_html = markdown(body_md, extensions=["tables", "fenced_code", "sane_lists", "nl2br"])
    soup = BeautifulSoup(body_html, "lxml")

    doc = Document()
    setup_docx_header_footer(doc)
    add_cover_docx(doc)

    # TOC
    toc_title = doc.add_paragraph()
    r = toc_title.add_run("Table of Contents")
    set_run_font(r, size=18, bold=True, color=(11, 58, 110))
    toc_p = doc.add_paragraph()
    add_toc_field(toc_p)

    # Manual TOC list (always visible without Word update)
    doc.add_paragraph()
    for title, _slug, level in build_toc_entries(body_md):
        p = doc.add_paragraph()
        indent = "    " if level > 1 else ""
        run = p.add_run(f"{indent}{title}")
        set_run_font(run, size=11 if level == 1 else 10, color=(17, 24, 39) if level == 1 else (75, 85, 99))

    doc.add_page_break()

    body = soup.body
    if not body:
        doc.save(OUT_DOCX)
        return

    for el in body.children:
        if getattr(el, "name", None) is None:
            continue
        name = el.name
        if name in ("h1", "h2", "h3"):
            level = int(name[1])
            p = doc.add_heading(el.get_text(strip=True), level=level)
            for run in p.runs:
                set_run_font(
                    run,
                    size={1: 16, 2: 13, 3: 12}[level],
                    bold=True,
                    color=(11, 58, 110) if level <= 2 else (17, 24, 39),
                )
        elif name == "p":
            p = doc.add_paragraph()
            html_inline_to_runs(p, el)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        elif name in ("ul", "ol"):
            for i, li in enumerate(el.find_all("li", recursive=False), 1):
                style = "List Number" if name == "ol" else "List Bullet"
                p = doc.add_paragraph(style=style)
                html_inline_to_runs(p, li)
        elif name == "table":
            rows = el.find_all("tr")
            if not rows:
                continue
            cols = max(len(r.find_all(["th", "td"])) for r in rows)
            table = doc.add_table(rows=len(rows), cols=cols)
            table.style = "Table Grid"
            for ri, row in enumerate(rows):
                cells = row.find_all(["th", "td"])
                for ci in range(cols):
                    cell = table.cell(ri, ci)
                    text = cells[ci].get_text(" ", strip=True) if ci < len(cells) else ""
                    cell.text = text
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            is_header = ri == 0 or (ci < len(cells) and cells[ci].name == "th")
                            set_run_font(
                                run,
                                size=9,
                                bold=is_header,
                                color=(255, 255, 255) if is_header and ri == 0 else (17, 24, 39),
                            )
                    if ri == 0:
                        shading = OxmlElement("w:shd")
                        shading.set(qn("w:fill"), "0B3A6E")
                        shading.set(qn("w:val"), "clear")
                        cell._tePr = cell._tc.get_or_add_tcPr()
                        cell._tc.get_or_add_tcPr().append(shading)
            doc.add_paragraph()
        elif name == "pre":
            p = doc.add_paragraph()
            run = p.add_run(el.get_text())
            set_run_font(run, name="Consolas", size=8, color=(15, 23, 42))
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
        elif name == "hr":
            doc.add_paragraph("─" * 40)
        elif name == "blockquote":
            p = doc.add_paragraph()
            run = p.add_run(el.get_text(" ", strip=True))
            set_run_font(run, size=10, color=(75, 85, 99))
            run.italic = True

    # Diagram note page
    doc.add_page_break()
    h = doc.add_heading("Visual Diagrams (Reference)", level=1)
    for run in h.runs:
        set_run_font(run, size=16, bold=True, color=(11, 58, 110))
    note = doc.add_paragraph()
    nr = note.add_run(
        "Architecture, deployment, roadmap, technology stack, and module hierarchy diagrams "
        "are rendered as vector graphics in the PDF edition of this proposal for print fidelity."
    )
    set_run_font(nr, size=11)

    doc.save(OUT_DOCX)


def pdf_page_count(path: Path) -> int:
    from pypdf import PdfReader

    return len(PdfReader(str(path)).pages)


def main() -> int:
    if not MD_PATH.exists():
        print(f"Missing source: {MD_PATH}", file=sys.stderr)
        return 1

    md_raw = MD_PATH.read_text(encoding="utf-8")
    html_doc = build_html(md_raw)
    print("Generating PDF...")
    generate_pdf(html_doc)
    print("Generating DOCX...")
    generate_docx(md_raw)

    pdf_size = OUT_PDF.stat().st_size
    docx_size = OUT_DOCX.stat().st_size
    pages = pdf_page_count(OUT_PDF)

    # Quick open verification
    pdf_ok = OUT_PDF.exists() and pdf_size > 10_000 and pages >= 5
    docx_ok = OUT_DOCX.exists() and docx_size > 5_000
    # Validate DOCX opens
    Document(str(OUT_DOCX))

    print("PDF_PATH=" + str(OUT_PDF))
    print("DOCX_PATH=" + str(OUT_DOCX))
    print(f"TOTAL_PAGES={pages}")
    print(f"PDF_SIZE_BYTES={pdf_size}")
    print(f"DOCX_SIZE_BYTES={docx_size}")
    print(f"VERIFICATION={'PASS' if pdf_ok and docx_ok else 'FAIL'}")
    return 0 if pdf_ok and docx_ok else 2


if __name__ == "__main__":
    raise SystemExit(main())
