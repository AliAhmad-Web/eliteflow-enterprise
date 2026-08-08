#!/usr/bin/env python3
"""
Generate EliteFlow — Internship Project Documentation (DOCX + PDF).

Layout matches Ali Ahmad Portfolio internship report style:
  cover → TOC → numbered chapters → centered figures → Word→PDF.

Run:
  python docs/scripts/generate_internship_report.py
"""

from __future__ import annotations

import json
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "docs"
ASSETS = OUT_DIR / "enterprise-manual-assets"
SHOTS = ASSETS / "screenshots-clean"
if not SHOTS.exists():
    SHOTS = ASSETS / "screenshots"
CATALOG_PATH = ASSETS / "screenshot-catalog.json"

DOCX_PATH = OUT_DIR / "ELITEFLOW_ENTERPRISE_DOCUMENTATION.docx"
PDF_PATH = OUT_DIR / "ELITEFLOW_ENTERPRISE_DOCUMENTATION.pdf"
PDF_FALLBACK = OUT_DIR / "ELITEFLOW_INTERNSHIP_DOCUMENTATION.pdf"

# Same visual system as Portfolio report (Calibri + accent + muted header/footer),
# with EliteFlow brand colors instead of cyan.
ACCENT = RGBColor(0x7C, 0x3A, 0xED)  # EliteFlow purple
DARK = RGBColor(0x0F, 0x17, 0x2A)
MUTED = RGBColor(0x47, 0x55, 0x69)
TABLE_HEADER_HEX = "0B3A6E"

WEB_URL = "https://eliteflow-web.vercel.app"
API_URL = "https://api-production-a778.up.railway.app"
GITHUB = "https://github.com/AliAhmad-Web/eliteflow-enterprise"

# (catalog page, figure number, caption) — figure numbers are document-local.
FIGURES: list[tuple[int, int, str]] = [
    (2, 1, "EliteFlow web login with email/password and social sign-in options."),
    (1, 2, "Google account picker continuing through Supabase OAuth."),
    (36, 3, "Forgot password form used to request a secure reset link."),
    (39, 4, "Forgot password confirmation after the reset email is queued."),
    (37, 5, "Reset password form with new password and confirmation fields."),
    (35, 6, "EliteFlow email verification / sign-in code message in Gmail."),
    (38, 7, "Password reset email delivered to the user’s inbox."),
    (4, 8, "Super admin console with tenants, health indicators, and actions."),
    (3, 9, "Client portal dashboard showing projects, invoices, and updates."),
    (5, 10, "Operations overview with revenue, clients, and project charts."),
    (6, 11, "Clients (CRM) list with contacts, emails, and status badges."),
    (7, 12, "Projects list with status, priority, and due dates."),
    (8, 13, "Tasks list with statuses, priorities, and assignees."),
    (9, 14, "Invoice & billing metrics with invoice list and statuses."),
    (10, 15, "Generate AI Document modal with document type and prompt."),
    (11, 16, "AI Documents library for generating and managing documents."),
    (12, 17, "AI Assistant chat for drafting emails and summaries."),
    (13, 18, "Team Messages with an Engineering group conversation open."),
    (14, 19, "Channels list for team, department, and group chats."),
    (16, 20, "Announcements feed with pinned system and hub updates."),
    (15, 21, "New announcement form with priority and expiration controls."),
    (18, 22, "Threads list for deployment and product discussions."),
    (17, 23, "New thread modal to start a discussion topic."),
    (20, 24, "Meetings list with scheduled sessions and AI summaries."),
    (19, 25, "Schedule meeting modal with times and waiting-room option."),
    (21, 26, "Activity feed of meetings, threads, and announcements."),
    (22, 27, "Calendar month view with meetings, deadlines, and events."),
    (23, 28, "File Manager library with upload and folder organization."),
    (24, 29, "Reports and analytics overview with KPIs and charts."),
    (25, 30, "Team / HR overview with workforce stats and leave approvals."),
    (26, 31, "Notification Center listing tasks, announcements, and alerts."),
    (27, 32, "Integration Center grid of third-party OAuth services."),
    (28, 33, "Security Center with sessions, login history, and score."),
    (29, 34, "Settings Center profile form with account and 2FA options."),
    (30, 35, "Downloads hub for desktop, Chrome extension, and Android builds."),
    (32, 36, "Chrome extension home with tasks and recent projects."),
    (31, 37, "Android sign-in screen with email and password fields."),
    (33, 38, "Android navigation drawer listing core EliteFlow modules."),
    (34, 39, "Android command center with KPIs and quick actions."),
]

PORTRAIT_PAGES = {5, 11, 19, 24, 27, 30, 31, 32, 33, 34}


def set_run_font(run, name="Calibri", size=11, bold=False, color=None, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr)
    run._r.append(fld_char_separate)
    run._r.append(fld_char_end)


def setup_header_footer(doc: Document):
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run("Ali Ahmad — EliteFlow · Internship Project Report")
    set_run_font(run, size=9, color=MUTED, italic=True)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = fp.add_run("Page ")
    set_run_font(r1, size=9, color=MUTED)
    add_page_number(fp)
    r2 = fp.add_run("  |  Internship Project Documentation")
    set_run_font(r2, size=9, color=MUTED)


def style_doc(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    for level, size in [(1, 18), (2, 14), (3, 12)]:
        style = styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style.font.color.rgb = DARK
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(14 if level == 1 else 10)
        style.paragraph_format.space_after = Pt(6)


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_para(
    doc,
    text,
    *,
    bold=False,
    italic=False,
    size=11,
    align="left",
    space_after=8,
    color=None,
):
    p = doc.add_paragraph()
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic, color=color)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            set_run_font(run, size=11)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Number")
        for run in p.runs:
            set_run_font(run, size=11)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                set_run_font(run, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), TABLE_HEADER_HEX)
        shading.set(qn("w:val"), "clear")
        hdr[i]._tc.get_or_add_tcPr().append(shading)
    for r_i, row in enumerate(rows):
        cells = table.rows[r_i + 1].cells
        for c_i, val in enumerate(row):
            cells[c_i].text = str(val)
            for p in cells[c_i].paragraphs:
                for run in p.runs:
                    set_run_font(run, size=10)
    doc.add_paragraph()


def shot_path(catalog_page: int) -> Path:
    return SHOTS / f"page-{catalog_page:02d}.png"


def add_figure(doc, catalog_page: int, fig_num: int, caption: str):
    path = shot_path(catalog_page)
    if not path.exists():
        add_para(doc, f"[Figure {fig_num} image missing: {path.name}]", italic=True, color=MUTED)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    if catalog_page in PORTRAIT_PAGES:
        run.add_picture(str(path), width=Inches(3.6))
    else:
        run.add_picture(str(path), width=Inches(6.1))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(f"Figure {fig_num}. {caption}")
    set_run_font(r, size=9, italic=True, color=MUTED)
    cap.paragraph_format.space_after = Pt(14)


def fig_by_num(fig_num: int) -> tuple[int, int, str]:
    for item in FIGURES:
        if item[1] == fig_num:
            return item
    raise KeyError(fig_num)


def insert_fig(doc, fig_num: int):
    catalog_page, num, caption = fig_by_num(fig_num)
    add_figure(doc, catalog_page, num, caption)


def page_break(doc):
    doc.add_page_break()


def build_cover(doc: Document):
    for _ in range(3):
        doc.add_paragraph()
    add_para(
        doc,
        "INTERNSHIP PROJECT DOCUMENTATION",
        align="center",
        size=12,
        bold=True,
        color=ACCENT,
        space_after=6,
    )
    add_para(
        doc,
        "EliteFlow – Enterprise Business Management System",
        align="center",
        size=24,
        bold=True,
        color=DARK,
        space_after=8,
    )
    add_para(
        doc,
        "Next.js · Express · Supabase · Electron · Expo · Chrome Extension · Vercel · Railway",
        align="center",
        size=11,
        italic=True,
        color=MUTED,
        space_after=28,
    )

    add_para(doc, "Submitted by", align="center", size=11, color=MUTED, space_after=2)
    add_para(doc, "Ali Ahmad", align="center", size=16, bold=True, space_after=2)
    add_para(doc, "Full-Stack Developer Intern", align="center", size=11, space_after=18)

    add_para(doc, "Project Brand", align="center", size=10, color=MUTED, space_after=2)
    add_para(doc, "EliteFlow", align="center", size=14, bold=True, color=ACCENT, space_after=18)

    add_para(doc, "Live Application", align="center", size=10, color=MUTED, space_after=2)
    add_para(doc, WEB_URL, align="center", size=11, space_after=6)
    add_para(doc, "API", align="center", size=10, color=MUTED, space_after=2)
    add_para(doc, API_URL, align="center", size=11, space_after=6)
    add_para(doc, "GitHub Repository", align="center", size=10, color=MUTED, space_after=2)
    add_para(doc, GITHUB, align="center", size=11, space_after=24)

    add_para(doc, f"Document Date: {date.today().strftime('%B %d, %Y')}", align="center", size=11, space_after=4)
    add_para(doc, "Document Version: 1.0", align="center", size=11, space_after=4)
    add_para(
        doc,
        "This report documents only features implemented in the repository and verified in production.",
        align="center",
        size=9,
        italic=True,
        color=MUTED,
    )
    page_break(doc)


def build_toc(doc: Document):
    add_heading(doc, "Table of Contents", 1)
    toc = [
        "1. Project Overview",
        "2. Internship Requirements Mapping",
        "3. Planning & Development Methodology",
        "4. System Architecture",
        "5. Authentication & Security",
        "6. Web Application",
        "7. Desktop Application (Electron)",
        "8. Android Application (Expo)",
        "9. Chrome Extension",
        "10. Backend API Design",
        "11. Database & Supabase",
        "12. AI Features",
        "13. Deployment & DevOps",
        "14. Testing & Verification",
        "15. Challenges & Solutions",
        "16. Future Improvements",
        "17. Conclusion",
        "18. References & Appendix",
    ]
    for item in toc:
        add_para(doc, item, size=12, space_after=6)
    page_break(doc)


def build_overview(doc: Document):
    add_heading(doc, "1. Project Overview", 1)
    add_para(
        doc,
        "EliteFlow is an enterprise business management platform built as an internship deliverable. "
        "It unifies CRM, projects, tasks, finance, collaboration, HR, analytics, and AI assistance "
        "into one product family spanning web, desktop, Android, and a Chrome extension. The web "
        "application is the primary console; other clients reuse the same Express API, Supabase Auth, "
        "and PostgreSQL data model.",
        align="justify",
    )

    add_heading(doc, "1.1 Project Objectives", 2)
    add_bullets(
        doc,
        [
            "Deliver a production-ready multi-tenant enterprise workspace for operations teams.",
            "Implement secure authentication (email/password, OAuth, password reset, verification).",
            "Enforce role-based access control across modules and API routes.",
            "Ship companion clients: Electron desktop, Expo Android, and Chrome MV3 extension.",
            "Integrate AI document generation and assistant workflows on top of business data.",
            "Deploy web on Vercel and API on Railway with Supabase for database, auth, and storage.",
        ],
    )

    add_heading(doc, "1.2 Scope of Work", 2)
    add_para(doc, "In scope (implemented):", bold=True)
    add_bullets(
        doc,
        [
            "Web modules: Dashboard, CRM, Projects, Tasks, Billing, Chat, Announcements, Threads, Meetings, Calendar, Files, Reports, Team/HR, Notifications, Integrations, Security, Settings.",
            "Auth flows: login, Google OAuth, forgot/reset password, email verification.",
            "Express REST API (/api/v1) with Prisma models against PostgreSQL.",
            "Desktop (Electron), Android (Expo), Chrome extension packages via Downloads Center.",
            "AI documents and AI assistant surfaces in the web product.",
        ],
    )
    add_para(doc, "Out of scope (not claimed in this report):", bold=True)
    add_bullets(
        doc,
        [
            "Full inventory ERP warehouse workflows (roadmap).",
            "Native iOS App Store release (Android/Expo focus for mobile).",
            "On-premise self-hosted appliance packaging.",
        ],
    )

    add_heading(doc, "1.3 Technology Stack", 2)
    add_table(
        doc,
        ["Layer", "Technology"],
        [
            ["Web", "Next.js (App Router), React, TypeScript, Tailwind CSS"],
            ["API", "Node.js, Express, Prisma, Zod validation"],
            ["Database / Auth / Storage", "Supabase (PostgreSQL + Auth + Storage)"],
            ["Desktop", "Electron wrapping the EliteFlow web shell"],
            ["Android", "Expo / React Native"],
            ["Extension", "Chrome Manifest V3"],
            ["AI", "Provider-backed assistant and document generation APIs"],
            ["Hosting", "Vercel (web), Railway (API), Supabase (data plane)"],
            ["Tooling", "Git/GitHub, ESLint, CI workflows"],
        ],
    )

    add_heading(doc, "1.4 Key URLs", 2)
    add_table(
        doc,
        ["Resource", "URL"],
        [
            ["Production web", WEB_URL],
            ["Production API", API_URL],
            ["GitHub", GITHUB],
        ],
    )
    page_break(doc)


def build_requirements(doc: Document):
    add_heading(doc, "2. Internship Requirements Mapping", 1)
    add_para(
        doc,
        "Internship work was delivered in incremental phases. The table below maps each required "
        "capability to the implemented artifact in this repository.",
        align="justify",
    )
    add_table(
        doc,
        ["Phase", "Requirement", "Implementation Status"],
        [
            ["Foundation", "Monorepo + Next.js web shell", "Completed — apps/web"],
            ["Foundation", "Express API + Prisma schema", "Completed — apps/api"],
            ["Foundation", "Supabase Auth + Postgres", "Completed"],
            ["Core Ops", "CRM / Projects / Tasks / Billing", "Completed — web modules + API"],
            ["Collaboration", "Chat, channels, announcements, threads", "Completed"],
            ["Collaboration", "Meetings + calendar", "Completed"],
            ["Intelligence", "AI documents + AI assistant", "Completed"],
            ["Clients", "Electron desktop package", "Completed — Downloads Center"],
            ["Clients", "Expo Android app", "Completed — Downloads Center"],
            ["Clients", "Chrome MV3 extension", "Completed — Downloads Center"],
            ["Hardening", "RBAC, sessions, security center", "Completed"],
            ["Release", "Vercel + Railway production deploy", "Completed"],
        ],
    )
    add_para(
        doc,
        "Screenshot evidence for authentication, operations modules, and companion clients appears "
        "in later chapters (Figures 1–39).",
        align="justify",
        italic=True,
        color=MUTED,
    )
    page_break(doc)


def build_planning(doc: Document):
    add_heading(doc, "3. Planning & Development Methodology", 1)
    add_heading(doc, "3.1 Approach", 2)
    add_para(
        doc,
        "Development followed an iterative, module-gated plan. Each module shipped behind shared auth "
        "and design-system foundations. AI tooling (Cursor) accelerated scaffolding; architecture, "
        "secrets, deployment, and acceptance testing remained human-owned.",
        align="justify",
    )
    add_heading(doc, "3.2 Work Breakdown", 2)
    add_numbered(
        doc,
        [
            "Establish monorepo, auth, and shell navigation.",
            "Ship core operational modules (CRM, projects, tasks, billing).",
            "Add collaboration (chat, announcements, threads, meetings).",
            "Add AI assistant and AI documents.",
            "Package desktop, Android, and Chrome extension clients.",
            "Harden security, settings, integrations; deploy and verify production.",
        ],
    )
    add_heading(doc, "3.3 Configuration Strategy", 2)
    add_bullets(
        doc,
        [
            "Secrets only in environment variables (Vercel / Railway / local .env), never hardcoded.",
            "Shared API client and RBAC helpers keep web and companion clients consistent.",
            "Prisma migrations applied before API rollout.",
        ],
    )
    page_break(doc)


def build_architecture(doc: Document):
    add_heading(doc, "4. System Architecture", 1)
    add_para(
        doc,
        "EliteFlow uses a multi-client, single-API architecture. Browser, Electron, Android, and the "
        "Chrome extension authenticate through Supabase and call the Express API. PostgreSQL stores "
        "tenant business data; Redis may back queues/cache where configured; object storage holds files.",
        align="justify",
    )
    add_heading(doc, "4.1 Logical Layers", 2)
    add_table(
        doc,
        ["Layer", "Responsibility"],
        [
            ["Presentation", "Next.js web, Electron shell, Expo Android, Chrome popup/panel"],
            ["API", "Express /api/v1 — validation, RBAC, business logic"],
            ["Data", "PostgreSQL via Prisma; Supabase Auth and Storage"],
            ["Edge / Hosting", "Vercel CDN for web; Railway for API process"],
        ],
    )
    add_heading(doc, "4.2 Deployment Topology", 2)
    add_para(
        doc,
        "Source control is GitHub. The web app deploys to Vercel. The API deploys to Railway. "
        "Supabase hosts PostgreSQL, Auth, and Storage. Desktop, Android, and extension builds are "
        "distributed through the in-app Downloads Center.",
        align="justify",
    )
    page_break(doc)


def build_auth(doc: Document):
    add_heading(doc, "5. Authentication & Security", 1)
    add_para(
        doc,
        "Authentication is powered by Supabase Auth. Users can sign in with email/password or Google "
        "OAuth. Password reset and email verification complete the account lifecycle. Session and "
        "device visibility is exposed in the Security Center for privileged users.",
        align="justify",
    )

    add_heading(doc, "5.1 Login", 2)
    add_para(
        doc,
        "The login screen is the entry point for web and companion clients that reuse the same auth "
        "backend. Social and password credentials both establish a session used for subsequent API calls.",
        align="justify",
    )
    insert_fig(doc, 1)
    insert_fig(doc, 2)

    add_heading(doc, "5.2 Password Recovery", 2)
    add_para(
        doc,
        "Forgot-password collects the account email, queues a reset message, and confirms delivery. "
        "The reset form accepts a new password with confirmation before updating credentials.",
        align="justify",
    )
    insert_fig(doc, 3)
    insert_fig(doc, 4)
    insert_fig(doc, 5)

    add_heading(doc, "5.3 Email Evidence", 2)
    add_para(
        doc,
        "Transactional messages (verification codes and password reset links) are delivered to the "
        "user inbox and verified during internship testing.",
        align="justify",
    )
    insert_fig(doc, 6)
    insert_fig(doc, 7)

    add_heading(doc, "5.4 Security Controls", 2)
    add_bullets(
        doc,
        [
            "HTTPS for all production endpoints.",
            "Role-based authorization on protected API routes.",
            "Session listing and security score surface in Security Center.",
            "Environment-separated secrets for Supabase, AI providers, and mail.",
        ],
    )
    page_break(doc)


def build_web(doc: Document):
    add_heading(doc, "6. Web Application", 1)
    add_para(
        doc,
        "The web application is the primary EliteFlow enterprise console. Built with Next.js App Router "
        "and React, it provides the reference UI for desktop, mobile, and extension experiences.",
        align="justify",
    )

    add_heading(doc, "6.1 Dashboards", 2)
    add_para(
        doc,
        "Role-aware dashboards surface KPIs and shortcuts. Super-admin console focuses on tenants and "
        "platform health; client portal emphasizes projects and invoices; operations overview presents "
        "revenue and delivery charts for internal staff.",
        align="justify",
    )
    insert_fig(doc, 8)
    insert_fig(doc, 9)
    insert_fig(doc, 10)

    add_heading(doc, "6.2 CRM, Projects, Tasks & Billing", 2)
    add_para(
        doc,
        "Clients maintain the commercial registry. Projects and tasks track delivery work with status "
        "and priority. Billing lists invoices with amounts and payment states — feeding finance KPIs.",
        align="justify",
    )
    insert_fig(doc, 11)
    insert_fig(doc, 12)
    insert_fig(doc, 13)
    insert_fig(doc, 14)

    add_heading(doc, "6.3 Collaboration", 2)
    add_para(
        doc,
        "Messages and channels support team conversation. Announcements broadcast org-wide updates. "
        "Threads capture longer-form product/ops discussions. Meetings and calendar coordinate time.",
        align="justify",
    )
    insert_fig(doc, 18)
    insert_fig(doc, 19)
    insert_fig(doc, 20)
    insert_fig(doc, 21)
    insert_fig(doc, 22)
    insert_fig(doc, 23)
    insert_fig(doc, 24)
    insert_fig(doc, 25)
    insert_fig(doc, 26)
    insert_fig(doc, 27)

    add_heading(doc, "6.4 Files, Reports & Team", 2)
    insert_fig(doc, 28)
    insert_fig(doc, 29)
    insert_fig(doc, 30)

    add_heading(doc, "6.5 Notifications, Integrations, Security & Settings", 2)
    insert_fig(doc, 31)
    insert_fig(doc, 32)
    insert_fig(doc, 33)
    insert_fig(doc, 34)

    add_heading(doc, "6.6 Downloads Center", 2)
    add_para(
        doc,
        "The Downloads hub distributes companion packages so users can install desktop, Android, and "
        "Chrome extension builds from the same authenticated workspace.",
        align="justify",
    )
    insert_fig(doc, 35)
    page_break(doc)


def build_desktop(doc: Document):
    add_heading(doc, "7. Desktop Application (Electron)", 1)
    add_para(
        doc,
        "The desktop client wraps the EliteFlow experience in Electron for always-on workstation use. "
        "It reuses the authenticated web shell and API contracts, reducing divergence while enabling "
        "native windowing and local install distribution via the Downloads Center.",
        align="justify",
    )
    add_heading(doc, "7.1 Features", 2)
    add_bullets(
        doc,
        [
            "Native window chrome for daily operations workflows.",
            "Same modules and RBAC as the web console.",
            "Distributed as a packaged installer from Downloads.",
        ],
    )
    add_heading(doc, "7.2 Architecture", 2)
    add_para(
        doc,
        "Electron main process hosts the shell; renderer loads the EliteFlow UI against the production "
        "API. Auth tokens follow the same Supabase session model as the browser.",
        align="justify",
    )
    add_heading(doc, "7.3 Benefits & Security", 2)
    add_bullets(
        doc,
        [
            "Faster access for power users without browser tab clutter.",
            "Inherits API authorization; no separate privilege model.",
            "Future scope: deeper OS integrations (tray, offline cache).",
        ],
    )
    page_break(doc)


def build_android(doc: Document):
    add_heading(doc, "8. Android Application (Expo)", 1)
    add_para(
        doc,
        "The Android application is built with Expo / React Native for on-the-go access to EliteFlow. "
        "It prioritizes sign-in, navigation, and a mobile command-center dashboard aligned with web KPIs.",
        align="justify",
    )
    insert_fig(doc, 37)
    insert_fig(doc, 38)
    insert_fig(doc, 39)

    add_heading(doc, "8.1 Features", 2)
    add_bullets(
        doc,
        [
            "Email/password sign-in against Supabase Auth.",
            "Drawer navigation across core modules.",
            "Mobile dashboard with KPIs and quick actions.",
        ],
    )
    add_heading(doc, "8.2 Architecture", 2)
    add_para(
        doc,
        "Expo app talks to the Express API over HTTPS. Builds can be produced with EAS and published "
        "through the Downloads Center for internship distribution.",
        align="justify",
    )
    add_heading(doc, "8.3 Benefits & Security", 2)
    add_bullets(
        doc,
        [
            "Field access to tasks, notifications, and KPIs.",
            "Secure token storage patterns appropriate for mobile.",
            "Future scope: push notifications and offline queues.",
        ],
    )
    page_break(doc)


def build_extension(doc: Document):
    add_heading(doc, "9. Chrome Extension", 1)
    add_para(
        doc,
        "The Chrome extension (Manifest V3) provides a lightweight popup for tasks and recent projects "
        "so users can act without opening the full web console.",
        align="justify",
    )
    insert_fig(doc, 36)

    add_heading(doc, "9.1 Features", 2)
    add_bullets(
        doc,
        [
            "Quick view of tasks and recent projects.",
            "Auth-aware session shared with EliteFlow account.",
            "Packaged MV3 build via Downloads Center.",
        ],
    )
    add_heading(doc, "9.2 Architecture", 2)
    add_para(
        doc,
        "Popup UI calls the same API endpoints used by web. Extension permissions are minimized to "
        "what the internship build requires.",
        align="justify",
    )
    add_heading(doc, "9.3 Benefits & Security", 2)
    add_bullets(
        doc,
        [
            "Reduces context switching for frequent check-ins.",
            "Uses HTTPS API + existing auth tokens.",
            "Future scope: page-context capture into tasks/CRM.",
        ],
    )
    page_break(doc)


def build_backend(doc: Document):
    add_heading(doc, "10. Backend API Design", 1)
    add_para(
        doc,
        "The backend is an Express application exposing versioned REST routes under /api/v1. Request "
        "validation uses Zod (or equivalent schema checks), and handlers enforce RBAC before mutating "
        "tenant data through Prisma.",
        align="justify",
    )
    add_heading(doc, "10.1 Representative Endpoints", 2)
    add_table(
        doc,
        ["Area", "Examples"],
        [
            ["Health", "GET /api/v1/health"],
            ["Clients", "CRUD /api/v1/clients"],
            ["Projects / Tasks", "/api/v1/projects, /api/v1/tasks"],
            ["Billing", "/api/v1/invoices"],
            ["Collaboration", "channels, messages, announcements, threads, meetings"],
            ["AI", "assistant + document generation routes"],
        ],
    )
    add_heading(doc, "10.2 Cross-cutting Concerns", 2)
    add_bullets(
        doc,
        [
            "CORS configured for the production web origin.",
            "Auth middleware validates bearer/session identity.",
            "Structured logging and health checks for Railway ops.",
        ],
    )
    page_break(doc)


def build_database(doc: Document):
    add_heading(doc, "11. Database & Supabase", 1)
    add_para(
        doc,
        "PostgreSQL is the system of record. Prisma models map domain entities (users/profiles, "
        "clients, projects, tasks, invoices, messages, meetings, etc.). Supabase provides Auth users, "
        "optional RLS policies where applicable, and Storage buckets for files.",
        align="justify",
    )
    add_heading(doc, "11.1 Data Domains", 2)
    add_bullets(
        doc,
        [
            "Identity & tenancy — users, roles, organizations.",
            "Commercial — clients, projects, invoices.",
            "Execution — tasks, meetings, calendar events.",
            "Collaboration — channels, messages, announcements, threads.",
            "Platform — notifications, integrations, security sessions, files.",
        ],
    )
    page_break(doc)


def build_ai(doc: Document):
    add_heading(doc, "12. AI Features", 1)
    add_para(
        doc,
        "EliteFlow includes an AI Documents library and an AI Assistant chat. Users generate structured "
        "documents from prompts and draft emails/summaries without leaving the workspace.",
        align="justify",
    )
    insert_fig(doc, 15)
    insert_fig(doc, 16)
    insert_fig(doc, 17)
    add_heading(doc, "12.1 Workflow", 2)
    add_numbered(
        doc,
        [
            "User selects document type or opens Assistant.",
            "Prompt and context are sent to the AI API route.",
            "Result is stored or inserted into the editor/chat transcript.",
            "User reviews, edits, and shares within EliteFlow modules.",
        ],
    )
    page_break(doc)


def build_deploy(doc: Document):
    add_heading(doc, "13. Deployment & DevOps", 1)
    add_para(
        doc,
        "Production deployment separates the static/SSR web tier from the long-running API process.",
        align="justify",
    )
    add_table(
        doc,
        ["Component", "Platform", "Notes"],
        [
            ["apps/web", "Vercel", WEB_URL],
            ["apps/api", "Railway", "Express /api/v1"],
            ["Database / Auth / Storage", "Supabase", "Postgres + Auth + Storage"],
            ["Desktop", "Electron builds", "Downloads Center"],
            ["Android", "Expo / EAS", "Downloads Center"],
            ["Extension", "Chrome MV3", "Downloads Center"],
        ],
    )
    add_para(
        doc,
        "Environment variables cover Supabase keys, AI provider secrets, CORS, and mail transport. "
        "Prisma migrations are applied before API rollout.",
        align="justify",
    )
    page_break(doc)


def build_testing(doc: Document):
    add_heading(doc, "14. Testing & Verification", 1)
    add_para(
        doc,
        "Verification combined manual UI walkthroughs with production smoke checks. Screenshots in "
        "this document are evidence of flows exercised on the live EliteFlow deployment.",
        align="justify",
    )
    add_table(
        doc,
        ["Area", "Verification"],
        [
            ["Auth", "Login, OAuth, forgot/reset, email receipt"],
            ["Core modules", "CRUD/list views for clients, projects, tasks, invoices"],
            ["Collaboration", "Messages, announcements, threads, meetings"],
            ["AI", "Document modal + assistant responses"],
            ["Clients", "Android screens + extension popup + downloads hub"],
            ["Deploy", "Web and API health on Vercel/Railway"],
        ],
    )
    page_break(doc)


def build_challenges(doc: Document):
    add_heading(doc, "15. Challenges & Solutions", 1)
    add_table(
        doc,
        ["Challenge", "Solution"],
        [
            ["Multi-client consistency", "Shared API contracts + Downloads Center distribution"],
            ["Auth across surfaces", "Centralize on Supabase Auth sessions"],
            ["Module sprawl", "Common shell, RBAC, and design system"],
            ["Production secrets", "Platform env vars; never commit credentials"],
            ["Screenshot evidence quality", "Clean captures for documentation figures"],
        ],
    )
    page_break(doc)


def build_future(doc: Document):
    add_heading(doc, "16. Future Improvements", 1)
    add_bullets(
        doc,
        [
            "Inventory module activation beyond roadmap placeholder.",
            "Deeper offline support for Android and desktop.",
            "Richer extension page-context actions.",
            "Expanded observability (APM, error budgets).",
            "Additional SSO providers for enterprise tenants.",
        ],
    )
    page_break(doc)


def build_conclusion(doc: Document):
    add_heading(doc, "17. Conclusion", 1)
    add_para(
        doc,
        "EliteFlow demonstrates a complete internship-scale enterprise product: authenticated multi-role "
        "web operations, AI-assisted productivity, and companion clients for desktop, Android, and Chrome. "
        "The system is deployed on Vercel, Railway, and Supabase, with features evidenced by the figures "
        "in this report.",
        align="justify",
    )
    add_para(
        doc,
        "The project prioritized working software over speculative claims — only implemented, tested "
        "capabilities are documented here.",
        align="justify",
    )
    page_break(doc)


def build_references(doc: Document):
    add_heading(doc, "18. References & Appendix", 1)
    add_heading(doc, "18.1 References", 2)
    add_bullets(
        doc,
        [
            "Next.js documentation — https://nextjs.org/docs",
            "Express documentation — https://expressjs.com/",
            "Supabase documentation — https://supabase.com/docs",
            "Electron documentation — https://www.electronjs.org/docs",
            "Expo documentation — https://docs.expo.dev/",
            "Chrome Extension MV3 — https://developer.chrome.com/docs/extensions/",
        ],
    )
    add_heading(doc, "18.2 Appendix A — Figure Index", 2)
    for _page, num, caption in FIGURES:
        add_para(doc, f"Figure {num}. {caption}", size=10, space_after=4)

    add_heading(doc, "18.3 Appendix B — Declaration", 2)
    add_para(
        doc,
        "I declare that this documentation describes the EliteFlow enterprise project as implemented "
        "in the GitHub repository and production deployment referenced herein. Features not present in "
        "the codebase have not been claimed.",
        align="justify",
    )
    add_para(
        doc,
        "AI tools (ChatGPT and Cursor AI) were used only for planning, documentation assistance, and "
        "implementation guidance. All code was reviewed, understood, tested, and integrated by the intern.",
        align="justify",
    )
    add_para(doc, "Intern: Ali Ahmad", space_after=2)
    add_para(doc, f"Date: {date.today().strftime('%B %d, %Y')}", space_after=2)


def build():
    if not SHOTS.exists():
        raise SystemExit(f"Screenshots folder missing: {SHOTS}")

    # Optional: validate catalog exists (not required for build).
    if CATALOG_PATH.exists():
        json.loads(CATALOG_PATH.read_text(encoding="utf-8"))

    doc = Document()
    style_doc(doc)
    setup_header_footer(doc)

    build_cover(doc)
    build_toc(doc)
    build_overview(doc)
    build_requirements(doc)
    build_planning(doc)
    build_architecture(doc)
    build_auth(doc)
    build_web(doc)
    build_desktop(doc)
    build_android(doc)
    build_extension(doc)
    build_backend(doc)
    build_database(doc)
    build_ai(doc)
    build_deploy(doc)
    build_testing(doc)
    build_challenges(doc)
    build_future(doc)
    build_conclusion(doc)
    build_references(doc)

    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    # If locked, write alternate name.
    try:
        doc.save(str(DOCX_PATH))
        out = DOCX_PATH
    except PermissionError:
        alt = OUT_DIR / "ELITEFLOW_ENTERPRISE_DOCUMENTATION_v2.docx"
        doc.save(str(alt))
        out = alt
        print(f"Primary DOCX locked; wrote {alt}")
    print(f"Wrote {out}")
    return out


def export_pdf_with_word(docx_path: Path, pdf_path: Path):
    import win32com.client  # type: ignore

    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    try:
        document = word.Documents.Open(str(docx_path.resolve()), ReadOnly=True)
        targets = [pdf_path, PDF_FALLBACK, OUT_DIR / "ELITEFLOW_ENTERPRISE_DOCUMENTATION_v2.pdf"]
        written = None
        last_err: Exception | None = None
        for target in targets:
            try:
                if target.exists():
                    try:
                        target.unlink()
                    except OSError:
                        # File locked (often open in a PDF viewer) — try next name.
                        continue
                document.SaveAs(str(target.resolve()), FileFormat=17)
                written = target
                break
            except Exception as exc:  # noqa: BLE001
                last_err = exc
        document.Close(False)
        if written is None:
            raise RuntimeError(f"Could not write PDF (last error: {last_err})")
        print(f"Wrote {written}")
        return written
    finally:
        try:
            word.Quit()
        except Exception:
            pass


def main():
    docx_path = build()
    try:
        import win32com.client  # noqa: F401
    except ImportError:
        import subprocess
        import sys

        subprocess.check_call([sys.executable, "-m", "pip", "install", "pywin32", "-q"])
    try:
        export_pdf_with_word(docx_path, PDF_PATH)
    except Exception as exc:
        print(f"PDF export via Word failed: {exc}")
        print("DOCX is available; convert manually if needed.")


if __name__ == "__main__":
    main()
